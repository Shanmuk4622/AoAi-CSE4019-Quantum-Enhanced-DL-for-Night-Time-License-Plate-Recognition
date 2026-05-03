import docx

def extract():
    doc = docx.Document('Quantum_LPR_Research_Paper.docx')
    with open('doc_content_with_tables.txt', 'w', encoding='utf-8') as f:
        for block in doc.element.body:
            if block.tag.endswith('p'):
                # Paragraph
                for p in doc.paragraphs:
                    if p._element == block:
                        f.write(p.text + '\n')
            elif block.tag.endswith('tbl'):
                # Table
                for table in doc.tables:
                    if table._element == block:
                        f.write('\n[TABLE START]\n')
                        for row in table.rows:
                            row_data = [cell.text.replace('\n', ' ').strip() for cell in row.cells]
                            f.write(' | '.join(row_data) + '\n')
                        f.write('[TABLE END]\n\n')

if __name__ == "__main__":
    extract()
